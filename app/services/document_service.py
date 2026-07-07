from io import BytesIO
from docx import Document
from app.models.job import Job
from app.services.storage_service import StorageService


class DocumentService:

    @staticmethod
    def generate_document(job: Job):

        document = Document()

        document.add_heading(
            "AI Generated Document",
            level=1
        )

        document.add_heading(
            "Original Request",
            level=2
        )

        document.add_paragraph(job.request)

        document.add_heading(
            "Generated Content",
            level=2
        )

        for item in job.generated_content:
            document.add_heading(
                item["title"],
                level=3
            )

            document.add_paragraph(
                item["output"]
            )

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        file_name = f"{job.id}.docx"

        public_url = (
            StorageService.upload_document(
                file_name,
                buffer.getvalue()
            )
        )

        return public_url